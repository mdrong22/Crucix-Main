"""
generateReviewDocx.py — RedLine Council Performance Review DOCX Generator
Called by generateReviewReport.mjs via child_process.execSync.

Usage:
    python3 scripts/generateReviewDocx.py '<stats_json>' '<output_path>'

Requires: python-docx  (pip install python-docx --break-system-packages)
"""

import sys
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Colour palette ───────────────────────────────────────────────────────────
C_DARK    = RGBColor(0x12, 0x12, 0x12)
C_PANEL   = RGBColor(0x1a, 0x1a, 0x1a)
C_ORANGE  = RGBColor(0xFF, 0x6B, 0x00)
C_GREEN   = RGBColor(0x00, 0xC8, 0x96)
C_RED     = RGBColor(0xFF, 0x44, 0x44)
C_YELLOW  = RGBColor(0xFF, 0xAA, 0x00)
C_LIGHT   = RGBColor(0xE0, 0xE0, 0xE0)
C_GRAY    = RGBColor(0x88, 0x88, 0x88)
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_HEADER  = RGBColor(0x0D, 0x0D, 0x0D)

FONT = "Courier New"

# ─── Helpers ──────────────────────────────────────────────────────────────────

def pct(v):
    if v is None: return "—"
    return f"{v*100:.1f}%"

def num(v, d=2):
    if v is None: return "—"
    return f"{v:.{d}f}"

def dollar(v):
    if v is None: return "—"
    return f"${abs(v):.0f}"

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), kwargs.get(edge, 'none'))
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', '333333'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def run(para, text, bold=False, color=None, size=11):
    r = para.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color or C_LIGHT
    return r

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    color = C_ORANGE
    size  = 16 if level == 1 else 12
    r = p.add_run(("◈ " if level == 1 else "▸ ") + text)
    r.font.name  = FONT
    r.font.size  = Pt(size)
    r.font.bold  = True
    r.font.color.rgb = color
    return p

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '333333')
    pBdr.append(bottom)
    pPr.append(pBdr)

def kpi_table(doc, kpis):
    """kpis: list of (label, value, color) tuples, displayed as a grid row."""
    table = doc.add_table(rows=2, cols=len(kpis))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    for i, (label, value, color) in enumerate(kpis):
        # Label row
        lc = table.cell(0, i)
        set_cell_bg(lc, '0D0D0D')
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(label)
        lr.font.name  = FONT
        lr.font.size  = Pt(8)
        lr.font.bold  = False
        lr.font.color.rgb = C_GRAY
        # Value row
        vc = table.cell(1, i)
        set_cell_bg(vc, '141414')
        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = vp.add_run(value)
        vr.font.name  = FONT
        vr.font.size  = Pt(16)
        vr.font.bold  = True
        vr.font.color.rgb = color or C_LIGHT
    doc.add_paragraph()

def data_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '0D0D0D')
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.name  = FONT
        r.font.size  = Pt(8)
        r.font.bold  = True
        r.font.color.rgb = C_GRAY
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, (text, color) in enumerate(row_data):
            cell = row.cells[i]
            set_cell_bg(cell, '111111')
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name  = FONT
            r.font.size  = Pt(10)
            r.font.color.rgb = color or C_LIGHT
    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

# ─── Main ─────────────────────────────────────────────────────────────────────

def generate(stats, date_label, output_path):
    doc = Document()

    # Page setup — US Letter, dark-ish
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(1)
    section.top_margin  = section.bottom_margin = Inches(1)

    # Default style
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(11)

    # ── Cover ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("REDLINE COUNCIL")
    r.font.name = FONT; r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = C_ORANGE
    p.add_run("\n")
    r2 = p.add_run("PERFORMANCE REVIEW")
    r2.font.name = FONT; r2.font.size = Pt(14); r2.font.bold = True; r2.font.color.rgb = C_LIGHT

    p2 = doc.add_paragraph()
    run(p2, f"Generated: {date_label}", color=C_GRAY, size=10)
    p2.add_run("   ")
    run(p2, f"Resolved: {stats['resolved']} / Total: {stats['totalDecisions']}", color=C_GRAY, size=10)

    divider(doc)
    doc.add_paragraph()

    # ── KPI Grid ──
    heading(doc, "EXECUTIVE SUMMARY")

    wr = stats['winRate']
    pf = stats['profitFactor']
    wr_color = C_GREEN if wr >= 0.55 else C_YELLOW if wr >= 0.40 else C_RED
    pf_color = C_GREEN if pf >= 1.5  else C_YELLOW if pf >= 1.0  else C_RED

    kpi_table(doc, [
        ("WIN RATE",      pct(stats['winRate']),      wr_color),
        ("PROFIT FACTOR", num(stats['profitFactor']), pf_color),
        ("AVG WIN",       f"+{num(stats['avgWinPct'])}%", C_GREEN),
        ("AVG LOSS",      f"{num(stats['avgLossPct'])}%", C_RED),
        ("WINS",          str(stats['wins']),          C_GREEN),
        ("LOSSES",        str(stats['losses']),        C_RED),
    ])

    doc.add_paragraph()

    # ── By Horizon ──
    heading(doc, "PERFORMANCE BY HORIZON", level=2)
    bh = stats.get('byHorizon', {})
    if bh:
        h_rows = []
        for h, s in bh.items():
            wr2 = s.get('winRate')
            ap  = s.get('avgPnlPct', 0)
            wr_c = C_GREEN if (wr2 or 0) >= 0.55 else C_YELLOW if (wr2 or 0) >= 0.40 else C_RED
            ap_c = C_GREEN if ap >= 0 else C_RED
            h_rows.append([
                (h,              C_LIGHT),
                (str(s['decisions']), C_LIGHT),
                (pct(wr2),       wr_c),
                (f"{ap:.2f}%",   ap_c),
            ])
        data_table(doc, ["HORIZON", "DECISIONS", "WIN RATE", "AVG P&L %"], h_rows, [1.5, 1.2, 1.2, 1.2])
    else:
        p = doc.add_paragraph()
        run(p, "No resolved decisions yet.", color=C_GRAY)

    # ── Signal Effectiveness ──
    heading(doc, "SIGNAL EFFECTIVENESS", level=2)
    bs = stats.get('bySignal', {})
    s_rows = []
    signal_labels = {
        'congressionalCluster': 'Congressional Cluster',
        'noCluster':            'No Cluster',
        'highSignalScore':      'Score ≥ 7',
        'lowSignalScore':       'Score < 7',
        'highVix':              'VIX ≥ 25',
    }
    for key, s in bs.items():
        wr3 = s.get('winRate')
        wr_c = C_GREEN if (wr3 or 0) >= 0.55 else C_YELLOW if (wr3 or 0) >= 0.40 else C_RED
        label = signal_labels.get(key, s.get('label', key))
        s_rows.append([
            (label,           C_LIGHT),
            (str(s['decisions']), C_LIGHT),
            (pct(wr3) if wr3 is not None else '—', wr_c),
        ])
    if s_rows:
        data_table(doc, ["SIGNAL", "DECISIONS", "WIN RATE"], s_rows, [2.8, 1.2, 1.2])

    # ── Top Winners ──
    heading(doc, "TOP WINNERS", level=2)
    tw = stats.get('topWins', [])
    if tw:
        w_rows = [[(d['ticker'], C_LIGHT), (d['horizon'], C_LIGHT),
                   (f"+{d['pnlPct']*100:.2f}%", C_GREEN), (dollar(d.get('pnlDollar')), C_GREEN)]
                  for d in tw]
        data_table(doc, ["TICKER", "HORIZON", "P&L %", "P&L $"], w_rows, [1.5, 1.3, 1.2, 1.2])
    else:
        p = doc.add_paragraph(); run(p, "No winners yet.", color=C_GRAY)

    # ── Top Losers ──
    heading(doc, "TOP LOSERS", level=2)
    tl = stats.get('topLosses', [])
    if tl:
        l_rows = [[(d['ticker'], C_LIGHT), (d['horizon'], C_LIGHT),
                   (f"{d['pnlPct']*100:.2f}%", C_RED), (f"-{dollar(d.get('pnlDollar'))}", C_RED)]
                  for d in tl]
        data_table(doc, ["TICKER", "HORIZON", "P&L %", "P&L $"], l_rows, [1.5, 1.3, 1.2, 1.2])
    else:
        p = doc.add_paragraph(); run(p, "No losses yet.", color=C_GRAY)

    # ── Recommendations ──
    doc.add_page_break()
    heading(doc, "STRATEGIC RECOMMENDATIONS")
    for rec in stats.get('recommendations', []):
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(rec)
        r.font.name = FONT
        r.font.size = Pt(11)
        r.font.color.rgb = C_LIGHT

    doc.add_paragraph()

    # ── Footer note ──
    divider(doc)
    p = doc.add_paragraph()
    run(p, "Generated by Crucix RedLine Review Council  ·  Data source: runs/decisions.json", color=C_GRAY, size=9)

    doc.save(output_path)
    print(f"[DocxGenerator] Saved: {output_path}")

# ─── Entry point ──────────────────────────────────────────────────────────────

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python3 generateReviewDocx.py '<json>' '<output_path>'")
        sys.exit(1)

    payload    = json.loads(sys.argv[1])
    stats      = payload['stats']
    date_label = payload.get('dateLabel', str(datetime.datetime.now()))
    out_path   = sys.argv[2]

    generate(stats, date_label, out_path)
